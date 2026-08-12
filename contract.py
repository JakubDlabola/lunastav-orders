import copy
import os
import re
import subprocess
import tempfile
from datetime import date

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'Smlouva-LUNASTAV-vzor.docx')

def _fmt_qty(q):
    try:
        q = float(q)
        return str(int(q)) if q == int(q) else str(q)
    except (TypeError, ValueError):
        return str(q) if q else ''

# Inline bold/highlight markers — control chars never present in normal text
_BOLD_ON       = '\x01B\x01'
_BOLD_OFF      = '\x01/B\x01'
_HIGHLIGHT_ON  = '\x01H\x01'
_HIGHLIGHT_OFF = '\x01/H\x01'
_MARKER_RE = re.compile(
    f'({re.escape(_BOLD_ON)}|{re.escape(_BOLD_OFF)}'
    f'|{re.escape(_HIGHLIGHT_ON)}|{re.escape(_HIGHLIGHT_OFF)})'
)


def fmt_czk(value):
    if not value:
        return '0'
    try:
        return '{:,.0f}'.format(float(value)).replace(',', '\xa0')
    except (TypeError, ValueError):
        return '0'


def fmt_date(value):
    if not value:
        return ''
    if isinstance(value, str):
        try:
            d = date.fromisoformat(value)
            return f'{d.day}. {d.month}. {d.year}'
        except Exception:
            return value
    return str(value)


def _make_bold(run):
    rPr = run.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.Element(qn('w:rPr'))
        run.insert(0, rPr)
    if rPr.find(qn('w:b')) is None:
        etree.SubElement(rPr, qn('w:b'))


def _make_highlight(run, color='yellow'):
    rPr = run.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.Element(qn('w:rPr'))
        run.insert(0, rPr)
    if rPr.find(qn('w:highlight')) is None:
        h = etree.SubElement(rPr, qn('w:highlight'))
        h.set(qn('w:val'), color)


def _build_run_text(r, text):
    """Append text (with \\t for tabs and \\n for line breaks) into run r."""
    lines = text.split('\n')
    for li, line in enumerate(lines):
        if li > 0:
            etree.SubElement(r, qn('w:br'))
        segments = line.split('\t')
        for i, seg in enumerate(segments):
            if seg:
                t = etree.SubElement(r, qn('w:t'))
                t.text = seg
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            if i < len(segments) - 1:
                etree.SubElement(r, qn('w:tab'))


def replace_in_element(element, replacements, bold_keys=None):
    """Replace placeholder text across all runs in every paragraph of an XML element."""
    for p in element.iter(qn('w:p')):
        runs = list(p.findall('.//' + qn('w:r')))
        if not runs:
            continue

        full_text = ''
        for r in runs:
            for child in r:
                if child.tag == qn('w:t'):
                    full_text += child.text or ''
                elif child.tag == qn('w:tab'):
                    full_text += '\t'

        if not any(k in full_text for k in replacements):
            continue

        should_bold = bool(bold_keys and any(k in full_text for k in bold_keys))

        new_text = full_text
        for k, v in replacements.items():
            new_text = new_text.replace(k, str(v))

        # Clear text/tab children from all runs
        first_r = runs[0]
        for r in runs:
            for child in list(r):
                if child.tag in (qn('w:t'), qn('w:tab')):
                    r.remove(child)

        if _BOLD_ON in new_text or _HIGHLIGHT_ON in new_text:
            parts = _MARKER_RE.split(new_text)
            parent = first_r.getparent()
            idx = list(parent).index(first_r)
            parent.remove(first_r)
            current_bold = should_bold
            current_highlight = False
            offset = 0
            for part in parts:
                if part == _BOLD_ON:
                    current_bold = True
                elif part == _BOLD_OFF:
                    current_bold = should_bold
                elif part == _HIGHLIGHT_ON:
                    current_highlight = True
                elif part == _HIGHLIGHT_OFF:
                    current_highlight = False
                elif part:
                    r = copy.deepcopy(first_r)
                    for child in list(r):
                        if child.tag in (qn('w:t'), qn('w:tab')):
                            r.remove(child)
                    if current_bold:
                        _make_bold(r)
                    if current_highlight:
                        _make_highlight(r)
                    _build_run_text(r, part)
                    parent.insert(idx + offset, r)
                    offset += 1
        else:
            if should_bold:
                _make_bold(first_r)
            _build_run_text(first_r, new_text)


def accept_track_changes(doc):
    """Accept all tracked changes: keep inserted text, discard deleted text."""
    body = doc.element.body

    # w:ins — accept insertions by unwrapping (keep children, remove wrapper)
    for ins in body.findall('.//' + qn('w:ins')):
        parent = ins.getparent()
        idx = list(parent).index(ins)
        for child in list(ins):
            parent.insert(idx, child)
            idx += 1
        parent.remove(ins)

    # w:del — reject deletions by removing entirely
    for del_el in body.findall('.//' + qn('w:del')):
        del_el.getparent().remove(del_el)

    # w:rPrChange / w:pPrChange — strip format-change markers
    for tag in (qn('w:rPrChange'), qn('w:pPrChange')):
        for el in body.findall('.//' + tag):
            el.getparent().remove(el)


def set_page_header(doc, contract_number: str):
    """Add contract number to header on all pages except the first."""
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    section = doc.sections[0]
    section.different_first_page_header_footer = True  # blank header on page 1

    header = section.header
    # Clear existing paragraphs
    for para in header.paragraphs:
        para.clear()

    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(contract_number)
    run.font.name = 'Arial'
    run.font.size = None  # inherit from style


def fill_items_table(doc, items):
    """Replace the single items template row with one row per order line."""
    table = doc.tables[1]
    template_tr = table.rows[2]._tr
    parent = template_tr.getparent()
    insert_idx = list(parent).index(template_tr)

    for i, item in enumerate(items):
        new_tr = copy.deepcopy(template_tr)
        replace_in_element(new_tr, {
            '{#items}': '',
            '{/items}': '',
            '{#hasItems}': '',
            '{/hasItems}': '',
            '{BusinessCaseItemCode}': item.get('code', ''),
            '{BusinessCaseItemName}': item.get('name', ''),
            '{BusinessCaseItemCount}': '1' if item.get('code') in ('D', '5000A', '5000B') else _fmt_qty(item.get('qty', 0)),
            '{BusinessCaseItemUnit}': 'ks' if item.get('code') in ('D', '5000A', '5000B') else item.get('unit', ''),
        })
        parent.insert(insert_idx + i, new_tr)

    parent.remove(template_tr)


def generate_contract(order: dict, partner: dict, lines: list) -> bytes:
    doc = Document(TEMPLATE_PATH)
    accept_track_changes(doc)
    set_page_header(doc, order.get('name', ''))

    TAX = 1.12
    LISTED_PRICES = {
        '3000A': 2002, '3100A': 2002, '3200A': 2002,
        '3000B': 751,  '3100B': 751,  '3200B': 751,
        '4000': 8000,
    }
    insul_codes = {'3000A', '3100A', '3200A', '3000B', '3100B', '3200B'}
    has_windows = False
    insul_area = 0.0
    real_listed_excl = 0.0
    doprava_subtotal = 0.0
    for ln in lines:
        if ln.get('display_type') or ln.get('is_downpayment'):
            continue
        code = re.sub(r'^\[(.+?)\].*', r'\1', ln['product_id'][1]) if isinstance(ln.get('product_id'), list) else ''
        qty = ln.get('product_uom_qty') or 0
        if code == '4000':
            has_windows = True
        if code in insul_codes:
            insul_area += qty
        if code in LISTED_PRICES:
            real_listed_excl += LISTED_PRICES[code] * qty / TAX
        elif code == 'D':
            doprava_subtotal += ln.get('price_subtotal') or 0
    dotace_area = '' if has_windows else (str(int(insul_area)) if insul_area else '')

    total_real_listed_excl = real_listed_excl + doprava_subtotal
    amount_untaxed = order.get('amount_untaxed') or 0
    real_discount_excl = max(0.0, total_real_listed_excl - amount_untaxed)
    discount_pct = max(3, round(real_discount_excl / total_real_listed_excl * 100)) if total_real_listed_excl else 3

    replacements = {
        '{code}':                 order.get('name', ''),
        '{companyName}':          partner.get('name', ''),
        '{companyBirthday}':      fmt_date(partner.get('x_studio_datum_narozeni') or ''),
        '{companyStreet}':        partner.get('street', '') or '',
        '{companyZipCode}':       partner.get('zip', '') or '',
        '{companyCity}':          partner.get('city', '') or '',
        '{companyEmail}':         partner.get('email', '') or '',
        '{companyTel1}':          partner.get('phone', '') or '',
        '{name}':                 _BOLD_ON + (lambda p: p.split('\n', 1)[0])(order.get('x_studio_popis_dila', '') or '') + _BOLD_OFF,
        '{Popis_dila_512bd}':     (lambda p: (_BOLD_ON + p.split('\n', 1)[1] + _BOLD_OFF) if '\n' in p else '')(order.get('x_studio_popis_dila', '') or ''),
        '{Adresa_rea_db304}':     _BOLD_ON + (
            order.get('x_studio_adresa_realizace') or ' '.join(
                p for p in [
                    partner.get('street', '') or '',
                    (partner.get('zip', '') or '') + ' ' + (partner.get('city', '') or ''),
                ] if p.strip()
            )
        ) + _BOLD_OFF,
        '{totalAmountWithTax}':   fmt_czk(order.get('amount_total')),
        '{totalAmount}':          fmt_czk(amount_untaxed),
        '{taxAmount}':            fmt_czk(order.get('amount_tax')),
        '{priceWithoutDiscount}': fmt_czk(total_real_listed_excl),
        '{discountPercent}':      str(discount_pct),
        '{discount}':             fmt_czk(real_discount_excl),
        '{_1_Zalohova_94eb7}':    _BOLD_ON + fmt_czk(order.get('x_studio_zaloha_kc')) + _BOLD_OFF,
        '{Termin_rea_c5929}':     _BOLD_ON + (order.get('x_studio_termin_zalohy_2') or '') + _BOLD_OFF,
        '{Doplatek_3d201}':       _BOLD_ON + fmt_czk(order.get('x_studio_doplatek_kc')) + _BOLD_OFF,
        '255286223/0600':         _BOLD_ON + '255286223/0600' + _BOLD_OFF,
        '{Platebni_p_0757d}':     _BOLD_ON + (order.get('x_studio_termin_dokonceni_2') or '') + _BOLD_OFF,
        '{Stavebni_p_5c162}':     _BOLD_ON + (order.get('x_studio_stavebni_pripravenost', '') or '') + _BOLD_OFF,
        '{scheduledEnd}':         fmt_date(order.get('x_studio_datum_podpisu_smlouvy')),
        '{Dotace_se__61533}':     dotace_area,
        '{Vyse_dotac_6d201}':     fmt_czk(order.get('x_studio_vyse_dotace_kc')) + '\xa0',
        '{Konecna_ce_0d59a}':     fmt_czk(order.get('x_studio_cena_po_odecteni_dotace')) + '\xa0',
        '{currency}':             'Kč',
        '{#hasItems}':            '',
        '{/hasItems}':            '',
        '{#items}':               '',
        '{/items}':               '',
    }

    # Handle items table first (row duplication), then replace everywhere else
    real_lines = [
        {
            'code': re.sub(r'^\[(.+?)\].*', r'\1', ln['product_id'][1])
                    if isinstance(ln.get('product_id'), list) else '',
            'name': re.sub(r'^\[.+?\]\s*', '', ln.get('name', '')),
            'qty': ln.get('product_uom_qty', 0),
            'unit': (ln.get('product_uom_id') or ln.get('product_uom') or [None, ''])[1],
        }
        for ln in lines
        if not ln.get('display_type') and not ln.get('is_downpayment')
    ]
    bold_keys = {
        '{totalAmountWithTax}', '{totalAmount}', '{taxAmount}',
        '{priceWithoutDiscount}', '{discount}',
        '{Vyse_dotac_6d201}', '{Konecna_ce_0d59a}',
    }
    fill_items_table(doc, real_lines)
    replace_in_element(doc.element.body, replacements, bold_keys=bold_keys)

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, 'contract.docx')
        pdf_path = os.path.join(tmpdir, 'contract.pdf')
        doc.save(docx_path)

        subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', tmpdir, docx_path],
            check=True,
            timeout=60,
            capture_output=True,
        )

        with open(pdf_path, 'rb') as f:
            return f.read()
